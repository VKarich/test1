from docx import Document
from datetime import datetime
from dateutil.relativedelta import *

INPUT_FILENAME = 'dkp.docx'
OUTPUT_FILENAME = "demo_in.docx"
WIDTH = 914400


class Tableter:
    @staticmethod
    def process_table(table, dates, payments):
        table.columns[0].width = WIDTH
        table.columns[1].width = WIDTH * 3
        table.columns[2].width = WIDTH * 2
        for i in range(len(dates)):
            key = str(i + 1)
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = "%s - %s" % (
                dates[key][0][:10],
                dates[key][1][:10]
            )
            row_cells[2].text = payments[key]

    @staticmethod
    def generate_file(model):
        document = Document(INPUT_FILENAME)
        if len(model["rsPaymentDate"]) > 0:
            Tableter.process_table(document.tables[0], model["rsPaymentDate"], model["rsPayment"])
        if len(model["rsPaymentDateFL"]) > 0:
            Tableter.process_table(document.tables[1], model["rsPaymentDateFL"], model["rsPaymentFL"])
        document.save(OUTPUT_FILENAME)